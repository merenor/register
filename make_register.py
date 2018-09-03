import os
from register import Register
import logging
import winsound
import docx
import countpages
import json

# shortcuts.py = dict with shortcuts
from shortcuts import shortcuts_lib

def main():

    logging.basicConfig(filename='logfile.log',level=logging.INFO,
        format='%(asctime)s %(message)s', datefmt='%d.%m.%Y-%H:%M:%S')

    r = Register()
    current_page = 0
    current_filename = ''

    print("Willkommen!")



    while True:

        user_input = input('\n[{}]: '.format(current_page))

        if user_input.startswith('.'):
            shortcut = user_input[1:]
            if shortcut in shortcuts_lib:
                print("Shortcut für ", shortcuts_lib[shortcut])
                user_input = shortcuts_lib[shortcut]
            else:
                print("Shortcut {} nicht gefunden.".format(shortcut))

        if len(user_input) > 1 and ' ' in user_input:
            command, parameters = user_input.split(' ', 1)
        else:
            command = user_input
            parameters = ''

######################

        if command == 'quit':
            r.close()
            logging.info("===== Bye. =====")
            break


        elif command == 'create':
            if parameters:
                r.create(parameters)
                print("Datei {} angelegt.".format(parameters))


        elif command == 'load':
            if parameters:
                r.load(parameters)
                print("Datei {} geladen.".format(parameters))
                logging.info("===== Datei {} geladen. =====".format(parameters))
                current_filename = parameters
            else:
                print("Bitte Dateinamen angeben.")


        elif command == 'showpagesfor':
            if parameters:
                pages_list = r.get_pages_by_name(parameters)
            str_pages = [str(p) for p in pages_list]
            print(", ".join(str_pages))


        elif command == 'find':
            if parameters:
                names_list = r.find_name(parameters)

                for entry in names_list:
                    print("({}): {}".format(entry[0], entry[1]))


        elif command == '+' or command == 'add':
            current_name = ''

            if parameters:
                names_list = r.find_name(parameters)

                if names_list:
                    print("Meintest du:")
                    if len(names_list) == 1: # only one entry
                        take = input("({}) {}\nÜbernehmen? (ENTER=ja, a=wie Eingabe anlegen) ".format(names_list[0][0], names_list[0][1]))
                        if take == '':
                            name_id = names_list[0][0]
                        elif take == 'a':
                            name_id = 'a'
                        else:
                            name_id = ' '
                    else:
                        for n in names_list:
                            print("({}): {}".format(n[0], n[1]))
                        name_id = input("id oder 'a' für wie Eingabe anlegen?: ")

                    if name_id == 'a':
                        current_name = parameters
                    else:
                        try:
                            name_id = int(name_id)
                        except:
                            print("Falsche id angegeben.")
                        for n in names_list:
                            if n[0] == name_id:
                                current_name = n[1]
                else:
                    add_new = input("Neu anlegen (j/n)? ")
                    if add_new == 'j':
                        current_name = parameters

            if current_name:
                if r.add_item(current_name, current_page):
                    print("'{}' auf Seite {} hinzugefügt.".format(current_name, current_page))
                    logging.info("add '{}' on '{}'".format(current_name, current_page))
                else:
                    print("Nichts getan.")

        elif command == '#':
            if parameters != '':
                try:
                    id = int(parameters)
                except:
                    print("ID muss Zahl sein.")
                    pass

            if r.add_item_by_id(id, current_page):
                entry = r.get_name_by_id(id)
                print("'{}' auf Seite '{}' hinzugefügt.".format(entry[1], current_page))
                logging.info("add '{}' on '{}'".format(entry[1], current_page))


        elif command == 'delfrompage':
            if parameters:
                names_list = r.find_name(parameters)

            if names_list:
                print("Meintest du:")
                for n in names_list:
                    print("({}): {}".format(n[0], n[1]))
                name_id = input("id oder 'n' für abbrechen?: ")

                # 'n' bricht ab
                if name_id == 'n':
                    pass
                else:
                    try:
                        name_id = int(name_id)
                    except:
                        print("Falsche ID angegeben.")
                        pass

                    entry = r.get_name_by_id(name_id)
                    found = False
                    for n in names_list:
                        if n[0] == entry[0]:
                            found = True
                            r.del_page(entry[0], current_page)
                            print("Name {} mit ID {} auf dieser Seite gelöscht.".format(entry[1], entry[0]))
                            logging.info("delfrompage '{}' on '{}'".format(entry[1], current_page))
                    if not found:
                        print("Name nicht gefunden. Richtige ID angegeben?")
            else:
                print("Keinen Namen gefunden.")



        elif command == 'delname':
            if parameters:
                names_list = r.find_name(parameters)

            if names_list:
                print("Meintest du:")
                for n in names_list:
                    print("({}): {}".format(n[0], n[1]))
                name_id = input("id oder 'n' für abbrechen?: ")

                # 'n' bricht ab
                if name_id == 'n':
                    pass
                else:
                    try:
                        name_id = int(name_id)
                    except:
                        print("Falsche ID angegeben.")
                        pass

                    for n in names_list:
                        if n[0] == name_id:
                            r.del_name(n[0])
                            print("Name {} mit ID {} gelöscht.".format(n[1], n[0]))
                            logging.info("delname '{}' id '{}'".format(n[1], n[0]))


        elif command == 'showall':
            names_list = r.get_all_names()

            for name in names_list:
                str_list = [str(p) for p in r.get_pages_by_name_id(name[0])]
                comma_list = ", ".join(str_list)
                print("({id}): {name}  {list}".format(id=name[0], name=name[1], list=comma_list))


        elif command == 'exporttoword':
            if parameters:
                doc = docx.Document()

                doc.add_heading("Register", 0)

                # get the list
                the_register = []
                names_list = r.get_all_names()

                for name in names_list:
                    #str_list = [str(p) for p in r.get_pages_by_name_id(name[0])]
                    comma_pages = countpages.countpages(r.get_pages_by_name_id(name[0]))
                    the_register.append({'name': name[1], 'pages': comma_pages})

                initial = 'A'
                doc.add_heading(initial, 1)

                for entry in the_register:
                    if entry['name'][0] != initial:
                        initial = entry['name'][0]
                        doc.add_heading(initial, 2)

                    # the 'voir' case: voir gets italic
                    if '. voir' in entry['name']:
                        a, b = entry['name'].split('. voir')
                        p = doc.add_paragraph(a + '. ')
                        p.add_run('voir').italic = True
                        p.add_run(b)
                        # no pages!
                    else:
                        doc.add_paragraph("{name}  {pages}".format(name=entry['name'], pages=entry['pages']))

                doc.save(parameters)


        elif command == 'exporttodict':
            if parameters:

                # get the list
                the_register = []
                names_list = r.get_all_names()

                for name in names_list:
                    the_register.append({'name': name[1], 'pages': r.get_pages_by_name_id(name[0])})

                with open(parameters, 'w', encoding='utf-8') as f:
                    f.write(json.dumps(the_register))


        elif command == 'here':
            # lists names on page, old 'namesonpage'
            names_list = r.get_names(current_page)

            for names in names_list:
                print("-", names)


        elif command == 'p' or command == 'page':
            if parameters != '':
                try:
                    new_page = int(parameters)
                    current_page = new_page
                    print("Seitenzahl auf {} geändert.".format(current_page))
                except:
                    print("Keine gültige Seitenzahl.")

        elif command == 'import':
            if parameters != '':
                r.import_names(parameters)


        elif command == 'rename':
            if parameters != '':
                try:
                    name_id = int(parameters)
                except:
                    print("Fehler: ID muss angegeben werden.")
                    pass

                new_name = input("Neuer Name: ")
                r.rename(name_id, new_name.strip())
                print("Geändert.")

        elif command == 'id':
            if parameters != '':
                try:
                    parameters = int(parameters)
                except:
                    print("ID muss eine Zahl sein.")
                    pass

                entry = r.get_name_by_id(int(parameters))
                if entry:
                    print("({id}) {name}".format(id=entry[0], name=entry[1]))
                else:
                    print("Nichts gefunden.")

        else:
            winsound.Beep(2000, 250)

if __name__ == '__main__':
    main()
