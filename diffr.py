import json
import sys
import docx

if len(sys.argv) != 4:
    exit("diffr datei_a.json date_b.json zieldatei.docx")


def load_jsonfile(filename):
    with open(filename) as f:
        return json.loads(f.read())


def get_pages_for(name, register):
    for e in register:
        if e['name'] == name:
            return e['pages']


register_a = load_jsonfile(sys.argv[1])
register_b = load_jsonfile(sys.argv[2])

reg_diff = []

for a in register_a:
    # if current name in a exists also in all names in b
    if a['name'] in [b['name'] for b in register_b]:
        pages_a = a['pages']
        pages_b = get_pages_for(a['name'], register_b)

        # find pages that are in a, but not in b
        page_diff = []
        for p in pages_a:
            if p not in pages_b:
                page_diff.append(p)
        if page_diff:
            reg_diff.append({'name': a['name'], 'pages': page_diff})
    # else append the whole data set
    else:
        reg_diff.append({'name': a['name'], 'pages': a['pages']})

reverse_reg_diff = []

for page in range(3, 466):
    names = []

    for entry in reg_diff:
        if page in entry['pages']:
            names.append(entry['name'])

    if names:
        reverse_reg_diff.append({'page': page, 'names': names})


# write doc
doc = docx.Document()

doc.add_heading("Namen in Datei '{}' und nicht in Datei '{}':".format(sys.argv[1], sys.argv[2]), 0)

for entry in reverse_reg_diff:
    doc.add_heading("S. " + str(entry['page']))
    for name in entry['names']:
        doc.add_paragraph("- " + str(name))

doc.save(sys.argv[3])

#doc.add_heading("Seiten in Datei '{}' und nicht in Datei '{}':".format(sys.argv[1], sys.argv[2]), 0)

#for entry in reg_diff:
#    doc.add_paragraph("{name}  {pages}".format(name=entry['name'], pages=entry['pages']))

#doc.save(sys.argv[3])
